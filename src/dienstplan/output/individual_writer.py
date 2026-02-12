"""Vereinfachtes Word-Dokument für individuelle Musiker-Dienstpläne.

Enthält nur:
1. Personalisierter Titel (Name, Position, Ensembles)
2. Wochenübersicht (Kalender)
3. Tagesansicht (Details)
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from ..models.plan import Dienstplan
from ..roster import Musician
from ..config import get_max_weekly_dienste
from .word_writer import (
    COLORS,
    _setup_document,
    _write_calendar_section,
    _write_detail_section,
    _add_section_break,
    _add_thin_line,
    _format_date_de,
)


def write_individual_docx(
    plan: Dienstplan,
    musician: Musician,
    output_path: str | Path,
    config: dict,
):
    """Schreibt einen individuellen Dienstplan als Word-Dokument."""
    doc = Document()
    max_weekly = get_max_weekly_dienste(config)

    _setup_document(doc)
    _write_individual_title(doc, plan, musician)
    _write_calendar_section(doc, plan, max_weekly)
    _add_section_break(doc)
    _write_detail_section(doc, plan)

    doc.save(str(output_path))


def _write_individual_title(doc: Document, plan: Dienstplan, musician: Musician):
    """Titel-Abschnitt für den individuellen Plan."""

    # Orchester-Name
    p = doc.add_paragraph()
    run = p.add_run("SÄCHSISCHE BLÄSERPHILHARMONIE")
    run.font.size = Pt(10)
    run.font.color.rgb = COLORS["secondary"]
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

    # "Dienstplan" + Musikername
    p = doc.add_paragraph()
    run = p.add_run(f"Dienstplan {musician.display_name}")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = COLORS["primary"]
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)

    # Position + Zusatz
    p = doc.add_paragraph()
    position_text = musician.position
    if musician.zusatz:
        position_text += f" — {musician.zusatz}"
    run = p.add_run(position_text)
    run.font.size = Pt(14)
    run.font.color.rgb = COLORS["secondary"]
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)

    # Zeitraum
    p = doc.add_paragraph()
    run = p.add_run(
        f"{_format_date_de(plan.plan_start)} – {_format_date_de(plan.plan_end)}"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = COLORS["secondary"]
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(4)

    # Ensemble-Mitgliedschaften
    if musician.ensembles:
        p = doc.add_paragraph()
        ens_text = "Ensembles: " + ", ".join(sorted(musician.ensembles))
        run = p.add_run(ens_text)
        run.font.size = Pt(10)
        run.font.color.rgb = COLORS["secondary"]
        run.font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(12)

    _add_thin_line(doc)
