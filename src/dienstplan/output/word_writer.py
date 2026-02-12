"""Word-Dokument-Ausgabe für den Dienstplan der SBP.

Modernes, klares Design ('Apple-Ästhetik') mit python-docx.
4 Abschnitte:
1. Wochen-Kalender mit TVK/HTV-Ampel
2. Tagesansicht (alle 10 Spalten)
3. TVK/HTV-Prüfbericht
4. Statistik-Übersicht
"""

from __future__ import annotations

from datetime import date, timedelta
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from ..models.plan import Dienstplan
from ..models.events import DienstType, Dienst
from ..constraints.base import Violation, Severity
from ..config import get_max_weekly_dienste, is_htv_active


_DE_MONTHS = [
    "", "Januar", "Februar", "März", "April", "Mai", "Juni",
    "Juli", "August", "September", "Oktober", "November", "Dezember",
]


def _format_date_de(d: date) -> str:
    """Deutsches Datumsformat: '01. März 2026'."""
    return f"{d.day:02d}. {_DE_MONTHS[d.month]} {d.year}"


# Farben (RGB-Tuples)
COLORS = {
    "primary": RGBColor(0x1D, 0x1D, 0x1F),      # Fast-Schwarz (Apple-Stil)
    "secondary": RGBColor(0x6E, 0x6E, 0x73),     # Grau
    "accent": RGBColor(0x00, 0x71, 0xE3),         # Apple-Blau
    "white": RGBColor(0xFF, 0xFF, 0xFF),
    "light_gray": RGBColor(0xF5, 0xF5, 0xF7),    # Hintergrund
    "border": RGBColor(0xD2, 0xD2, 0xD7),         # Rahmen
    "frei": RGBColor(0x34, 0xC7, 0x59),           # Grün
    "probe": RGBColor(0x00, 0x71, 0xE3),          # Blau
    "konzert": RGBColor(0xFF, 0x37, 0x5F),        # Rot/Rosa
    "tour": RGBColor(0xFF, 0x9F, 0x0A),           # Orange
    "sonstiges": RGBColor(0x8E, 0x8E, 0x93),      # Grau
    "error": RGBColor(0xFF, 0x37, 0x5F),           # Rot
    "warning": RGBColor(0xFF, 0x9F, 0x0A),         # Orange
    "ok": RGBColor(0x34, 0xC7, 0x59),              # Grün
}

DIENST_TYPE_COLORS = {
    DienstType.FREI: COLORS["frei"],
    DienstType.URLAUB: COLORS["frei"],
    DienstType.REISEZEITAUSGLEICH: COLORS["secondary"],
    DienstType.PROBE: COLORS["probe"],
    DienstType.GENERALPROBE: COLORS["probe"],
    DienstType.HAUPTPROBE: COLORS["probe"],
    DienstType.ANSPIELPROBE: COLORS["probe"],
    DienstType.KONZERT: COLORS["konzert"],
    DienstType.ABO_KONZERT: COLORS["konzert"],
    DienstType.GASTSPIEL: COLORS["konzert"],
    DienstType.SCHUELERKONZERT: COLORS["tour"],
    DienstType.BABYKONZERT: COLORS["tour"],
    DienstType.DIRIGIERKURS: COLORS["tour"],
    DienstType.REISE: COLORS["tour"],
    DienstType.SONSTIGES: COLORS["sonstiges"],
}


def write_dienstplan_docx(plan: Dienstplan, output_path: str | Path, config: dict):
    """Schreibt den Dienstplan als modernes Word-Dokument."""
    doc = Document()
    max_weekly = get_max_weekly_dienste(config)

    _setup_document(doc)
    _write_title(doc, plan)
    _write_calendar_section(doc, plan, max_weekly)
    _add_section_break(doc)
    _write_detail_section(doc, plan)
    _add_section_break(doc)
    _write_violations_section(doc, plan)
    _add_section_break(doc)
    _write_statistics_section(doc, plan, max_weekly)
    _add_section_break(doc)
    _write_ausgleichszeitraum_section(doc, plan, config)

    doc.save(str(output_path))


def _setup_document(doc: Document):
    """Setzt Grundformatierung des Dokuments."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = COLORS["primary"]

    # Seitenränder reduzieren
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)


def _write_title(doc: Document, plan: Dienstplan):
    """Schreibt die Titelzeile."""
    # Orchester-Name
    p = doc.add_paragraph()
    run = p.add_run("SÄCHSISCHE BLÄSERPHILHARMONIE")
    run.font.size = Pt(10)
    run.font.color.rgb = COLORS["secondary"]
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

    # Dienstplan-Titel
    p = doc.add_paragraph()
    run = p.add_run("Dienstplan")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLORS["primary"]
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)

    # Zeitraum
    p = doc.add_paragraph()
    run = p.add_run(
        f"{_format_date_de(plan.plan_start)} – {_format_date_de(plan.plan_end)}"
    )
    run.font.size = Pt(14)
    run.font.color.rgb = COLORS["secondary"]
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(20)

    # Trennlinie
    _add_thin_line(doc)


def _write_calendar_section(doc: Document, plan: Dienstplan, max_weekly: int):
    """Abschnitt 1: Wochen-Kalenderansicht."""
    _add_section_heading(doc, "1", "Wochenübersicht")

    # Tabelle: KW | Mo | Di | Mi | Do | Fr | Sa | So | Σ | Status
    table = doc.add_table(rows=1, cols=10)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    headers = ["KW", "Mo", "Di", "Mi", "Do", "Fr", "Sa", "So", "Σ", ""]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = text
        _style_header_cell(cell)

    # Wochen-Zeilen
    for week in plan.weeks:
        row = table.add_row()

        # KW
        row.cells[0].text = str(week.week_number)
        _style_body_cell(row.cells[0], bold=True, align="center", font_size=9)

        # Tage Mo-So
        for day_offset in range(7):
            day_date = week.start_date + timedelta(days=day_offset)
            cell = row.cells[day_offset + 1]
            dienst = week.dienst_for_date(day_date)

            if dienst:
                lines = [day_date.strftime("%d.%m.")]
                if dienst.is_free:
                    type_label = _get_free_label(dienst)
                    lines.append(type_label)
                elif dienst.dienst_count > 0:
                    for e in dienst.events[:2]:  # Max 2 Events anzeigen
                        label = e.dienst_type.value
                        if e.start_time:
                            label += f" {e.start_time.strftime('%H:%M')}"
                            if e.end_time:
                                label += f"–{e.end_time.strftime('%H:%M')}"
                        lines.append(label)
                    lines.append(f"[{dienst.dienst_count:g}]")

                cell.text = ""
                p = cell.paragraphs[0]
                # Datum
                run = p.add_run(lines[0] + "\n")
                run.font.size = Pt(7)
                run.font.color.rgb = COLORS["secondary"]
                run.font.name = 'Calibri'
                # Events
                for line in lines[1:]:
                    run = p.add_run(line + "\n")
                    run.font.size = Pt(7)
                    run.font.name = 'Calibri'
                    if "[" in line:
                        run.font.bold = True
                        run.font.color.rgb = COLORS["primary"]
                    elif dienst.is_free:
                        run.font.color.rgb = COLORS["frei"]
                    else:
                        color = DIENST_TYPE_COLORS.get(dienst.primary_type, COLORS["primary"])
                        run.font.color.rgb = color

                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)

                # Hintergrundfarbe für freie Tage
                if dienst.is_free:
                    _set_cell_bg(cell, "E8F8EE")  # Zartes Grün
            else:
                cell.text = day_date.strftime("%d.%m.")
                _style_body_cell(cell, font_size=7, color=COLORS["secondary"])

            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        # Dienste-Summe
        total = week.total_dienste
        row.cells[8].text = f"{total:g}"
        _style_body_cell(row.cells[8], bold=True, align="center")

        # TVK/HTV-Ampel
        status = week.tvk_status_for(max_weekly)
        status_cell = row.cells[9]
        if status == "error":
            status_cell.text = "!!!"
            _style_body_cell(status_cell, bold=True, align="center")
            _set_cell_bg(status_cell, "FFE5EA")
            status_cell.paragraphs[0].runs[0].font.color.rgb = COLORS["error"]
        elif status == "warning":
            status_cell.text = "!"
            _style_body_cell(status_cell, bold=True, align="center")
            _set_cell_bg(status_cell, "FFF3E0")
            status_cell.paragraphs[0].runs[0].font.color.rgb = COLORS["warning"]
        else:
            status_cell.text = "OK"
            _style_body_cell(status_cell, align="center")
            _set_cell_bg(status_cell, "E8F8EE")
            status_cell.paragraphs[0].runs[0].font.color.rgb = COLORS["ok"]

    # Spaltenbreiten setzen
    _set_column_widths(table, [0.8, 2.0, 2.0, 2.0, 2.0, 2.0, 2.0, 2.0, 0.8, 0.6])

    # Zusammenfassung
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(f"Gesamt: {plan.total_dienste:g} Dienste in {plan.total_weeks} Wochen")
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = COLORS["primary"]
    run.font.name = 'Calibri'


def _write_detail_section(doc: Document, plan: Dienstplan):
    """Abschnitt 2: Tagesansicht mit allen 10 Spalten."""
    _add_section_heading(doc, "2", "Tagesansicht")

    # Tabelle: Tag | Datum | Zeit | Formation | Ort | Programm | Leitung | Kleidung | Dienste | Sonstiges
    headers = ["Tag", "Datum", "Zeit", "Formation", "Ort", "Programm", "Leitung", "Kleidung", "D.", "Sonstiges"]
    table = doc.add_table(rows=1, cols=10)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, text in enumerate(headers):
        _style_header_cell(table.rows[0].cells[i])
        table.rows[0].cells[i].text = text

    row_idx = 0
    for dienst in plan.all_dienste():
        if not dienst.events and dienst.dienst_count == 0:
            # Freier Tag
            row = table.add_row()
            row_idx += 1
            row.cells[0].text = dienst.day_of_week
            row.cells[1].text = dienst.dienst_date.strftime("%d.%m.%Y")
            row.cells[5].text = "Frei" if dienst.is_free else ""
            row.cells[8].text = "0"
            for j in range(10):
                _style_body_cell(row.cells[j], font_size=8)
                if dienst.is_free and row.cells[j].paragraphs[0].runs:
                    row.cells[j].paragraphs[0].runs[0].font.color.rgb = COLORS["frei"]
            if row_idx % 2 == 0:
                _shade_row(row, "F8F9FA")
        else:
            for event in dienst.events:
                row = table.add_row()
                row_idx += 1

                zeit = ""
                if event.start_time:
                    zeit = event.start_time.strftime("%H:%M")
                    if event.end_time:
                        zeit += f"–{event.end_time.strftime('%H:%M')}"

                row.cells[0].text = dienst.day_of_week
                row.cells[1].text = dienst.dienst_date.strftime("%d.%m.%Y")
                row.cells[2].text = zeit
                row.cells[3].text = event.formation.value if event.formation else ""
                row.cells[4].text = event.ort
                row.cells[5].text = event.programm
                row.cells[6].text = event.leitung
                row.cells[7].text = event.kleidung
                row.cells[8].text = f"{dienst.dienst_count:g}"
                row.cells[9].text = event.sonstiges

                for j in range(10):
                    _style_body_cell(row.cells[j], font_size=8)

                if row_idx % 2 == 0:
                    _shade_row(row, "F8F9FA")

            if not dienst.events:
                row = table.add_row()
                row_idx += 1
                row.cells[0].text = dienst.day_of_week
                row.cells[1].text = dienst.dienst_date.strftime("%d.%m.%Y")
                row.cells[8].text = f"{dienst.dienst_count:g}"
                for j in range(10):
                    _style_body_cell(row.cells[j], font_size=8)
                if row_idx % 2 == 0:
                    _shade_row(row, "F8F9FA")

    _set_column_widths(table, [0.7, 1.5, 1.3, 1.5, 2.0, 3.5, 1.2, 2.2, 0.5, 2.5])


def _write_violations_section(doc: Document, plan: Dienstplan):
    """Abschnitt 3: TVK/HTV-Prüfbericht."""
    _add_section_heading(doc, "3", "TVK/HTV-Prüfbericht")

    if not plan.violations:
        p = doc.add_paragraph()
        run = p.add_run("Keine Verstöße gefunden.")
        run.font.color.rgb = COLORS["ok"]
        run.font.italic = True
        run.font.name = 'Calibri'
        return

    # Zusammenfassung
    errors = sum(1 for v in plan.violations if v.severity == Severity.ERROR)
    warnings = sum(1 for v in plan.violations if v.severity == Severity.WARNING)
    infos = sum(1 for v in plan.violations if v.severity == Severity.INFO)

    p = doc.add_paragraph()
    if errors:
        run = p.add_run(f"{errors} Fehler  ")
        run.font.color.rgb = COLORS["error"]
        run.font.bold = True
        run.font.name = 'Calibri'
    if warnings:
        run = p.add_run(f"{warnings} Warnungen  ")
        run.font.color.rgb = COLORS["warning"]
        run.font.bold = True
        run.font.name = 'Calibri'
    if infos:
        run = p.add_run(f"{infos} Hinweise")
        run.font.color.rgb = COLORS["secondary"]
        run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(10)

    # Tabelle
    headers = ["", "Regel", "§", "KW/Datum", "Ist", "Limit", "Beschreibung"]
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, text in enumerate(headers):
        _style_header_cell(table.rows[0].cells[i])
        table.rows[0].cells[i].text = text

    for v in plan.violations:
        row = table.add_row()

        # Schwere-Icon
        if v.severity == Severity.ERROR:
            row.cells[0].text = "!!!"
            _set_cell_bg(row.cells[0], "FFE5EA")
        elif v.severity == Severity.WARNING:
            row.cells[0].text = "!"
            _set_cell_bg(row.cells[0], "FFF3E0")
        else:
            row.cells[0].text = "i"
            _set_cell_bg(row.cells[0], "E3F2FD")

        row.cells[1].text = v.rule_name
        row.cells[2].text = v.tvk_paragraph or ""

        if v.affected_week:
            row.cells[3].text = f"KW {v.affected_week}"
        elif v.affected_dates:
            row.cells[3].text = v.affected_dates[0].strftime("%d.%m.")

        row.cells[4].text = str(v.current_value) if v.current_value is not None else ""
        row.cells[5].text = str(v.limit_value) if v.limit_value is not None else ""
        row.cells[6].text = v.message

        for j in range(7):
            _style_body_cell(row.cells[j], font_size=8)

    _set_column_widths(table, [0.5, 2.0, 1.0, 1.0, 0.7, 0.7, 5.5])


def _write_statistics_section(doc: Document, plan: Dienstplan, max_weekly: int):
    """Abschnitt 4: Statistik-Dashboard."""
    _add_section_heading(doc, "4", "Statistik")

    # Kennzahlen als saubere Liste
    stats = [
        ("Zeitraum", f"{plan.plan_start.strftime('%d.%m.%Y')} – {plan.plan_end.strftime('%d.%m.%Y')}"),
        ("Kalenderwochen", str(plan.total_weeks)),
        ("Gesamt-Dienste", f"{plan.total_dienste:g}"),
        ("Durchschnitt/Woche", f"{plan.avg_dienste_per_week}"),
        ("Freie Tage", str(plan.total_free_days)),
        ("Freie Sonntage", str(plan.free_sundays)),
        ("Max. Dienste/Woche (HTV)", str(max_weekly)),
    ]

    table = doc.add_table(rows=len(stats), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (label, value) in enumerate(stats):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        _style_body_cell(table.rows[i].cells[0], bold=True, font_size=10)
        _style_body_cell(table.rows[i].cells[1], font_size=10)

    _set_column_widths(table, [4.0, 6.0])

    # Dienste pro Woche
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Dienste pro Woche")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = COLORS["primary"]
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(6)

    week_table = doc.add_table(rows=1, cols=3)
    week_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(["KW", "Dienste", "Status"]):
        _style_header_cell(week_table.rows[0].cells[i])
        week_table.rows[0].cells[i].text = h

    for week in plan.weeks:
        row = week_table.add_row()
        row.cells[0].text = str(week.week_number)
        row.cells[1].text = f"{week.total_dienste:g}"

        status = week.tvk_status_for(max_weekly)
        if status == "error":
            row.cells[2].text = "Verstoß"
        elif status == "warning":
            row.cells[2].text = "Am Limit"
        else:
            row.cells[2].text = "OK"

        for j in range(3):
            _style_body_cell(row.cells[j], font_size=9)
        if status == "error":
            row.cells[1].paragraphs[0].runs[0].font.color.rgb = COLORS["error"]
            row.cells[1].paragraphs[0].runs[0].font.bold = True
            row.cells[2].paragraphs[0].runs[0].font.color.rgb = COLORS["error"]
        elif status == "warning":
            row.cells[1].paragraphs[0].runs[0].font.color.rgb = COLORS["warning"]
            row.cells[2].paragraphs[0].runs[0].font.color.rgb = COLORS["warning"]
        else:
            row.cells[2].paragraphs[0].runs[0].font.color.rgb = COLORS["ok"]

    _set_column_widths(week_table, [2.0, 3.0, 3.0])

    # Dienste nach Typ
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Dienste nach Typ")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = COLORS["primary"]
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(6)

    by_type = plan.dienste_by_type()
    if by_type:
        type_table = doc.add_table(rows=1, cols=2)
        type_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, h in enumerate(["Typ", "Anzahl Tage"]):
            _style_header_cell(type_table.rows[0].cells[i])
            type_table.rows[0].cells[i].text = h

        for dtype, count in sorted(by_type.items(), key=lambda x: -x[1]):
            row = type_table.add_row()
            row.cells[0].text = dtype.value
            row.cells[1].text = str(count)
            for j in range(2):
                _style_body_cell(row.cells[j], font_size=9)

        _set_column_widths(type_table, [4.0, 3.0])


def _write_ausgleichszeitraum_section(doc: Document, plan: Dienstplan, config: dict):
    """Abschnitt 5: HTV §12,2 Ausgleichszeitraum-Auswertung.

    Zeigt die festen 24-Wochen-Perioden mit Dienste-Summen.
    Wird immer ausgegeben (Audit-Pflicht laut HTV).
    """
    _add_section_heading(doc, "5", "Ausgleichszeitraum (HTV §12,2)")

    if not is_htv_active(config):
        p = doc.add_paragraph()
        run = p.add_run("HTV nicht aktiv – Ausgleichszeitraum nicht relevant.")
        run.font.color.rgb = COLORS["secondary"]
        run.font.italic = True
        run.font.name = 'Calibri'
        return

    azr = config.get("htv", {}).get("ausgleichszeitraum", {})
    window_size = azr.get("weeks", 24)
    max_dienste = azr.get("max_dienste", 183)
    transfer = azr.get("transfer_dienste", 9)
    period_1_start_str = azr.get("period_1_start", "2026-08-17")

    from datetime import date as date_cls
    period_1_start = date_cls.fromisoformat(period_1_start_str)
    period_1_end = period_1_start + timedelta(weeks=window_size) - timedelta(days=1)
    period_2_start = period_1_end + timedelta(days=1)
    period_2_end = period_2_start + timedelta(weeks=window_size) - timedelta(days=1)

    periods = [
        ("Periode 1", period_1_start, period_1_end),
        ("Periode 2", period_2_start, period_2_end),
    ]

    # Erklärungstext
    p = doc.add_paragraph()
    run = p.add_run(
        f"Max. {max_dienste} Dienste pro {window_size}-Wochen-Zeitraum. "
        f"Bis zu {transfer} Dienste zwischen Perioden übertragbar."
    )
    run.font.size = Pt(10)
    run.font.color.rgb = COLORS["secondary"]
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(10)

    # Tabelle
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(["Periode", "Zeitraum", "Dienste", "Limit", "Status"]):
        _style_header_cell(table.rows[0].cells[i])
        table.rows[0].cells[i].text = h

    for label, p_start, p_end in periods:
        total = 0.0
        matching_weeks = []
        for week in plan.weeks:
            if week.end_date >= p_start and week.start_date <= p_end:
                total += week.total_dienste
                matching_weeks.append(week)

        row = table.add_row()
        row.cells[0].text = label

        if matching_weeks:
            first_kw = matching_weeks[0].week_number
            last_kw = matching_weeks[-1].week_number
            weeks_count = len(matching_weeks)
            row.cells[1].text = (
                f"KW {first_kw}–{last_kw}\n"
                f"({p_start.strftime('%d.%m.%Y')} – {p_end.strftime('%d.%m.%Y')})\n"
                f"{weeks_count} von {window_size} Wochen im Plan"
            )
            row.cells[2].text = f"{total:g}"
            row.cells[3].text = str(max_dienste)

            if total > max_dienste:
                row.cells[4].text = f"ÜBERSCHRITTEN (+{total - max_dienste:g})"
                for j in range(5):
                    _style_body_cell(row.cells[j], font_size=9)
                row.cells[4].paragraphs[0].runs[0].font.color.rgb = COLORS["error"]
                row.cells[4].paragraphs[0].runs[0].font.bold = True
                row.cells[2].paragraphs[0].runs[0].font.color.rgb = COLORS["error"]
                row.cells[2].paragraphs[0].runs[0].font.bold = True
            else:
                puffer = max_dienste - total
                row.cells[4].text = f"OK ({puffer:g} Puffer)"
                for j in range(5):
                    _style_body_cell(row.cells[j], font_size=9)
                row.cells[4].paragraphs[0].runs[0].font.color.rgb = COLORS["ok"]
        else:
            row.cells[1].text = (
                f"{p_start.strftime('%d.%m.%Y')} – {p_end.strftime('%d.%m.%Y')}\n"
                f"Keine Daten im Plan"
            )
            row.cells[2].text = "–"
            row.cells[3].text = str(max_dienste)
            row.cells[4].text = "Kein Daten"
            for j in range(5):
                _style_body_cell(row.cells[j], font_size=9)
            row.cells[4].paragraphs[0].runs[0].font.color.rgb = COLORS["secondary"]

    _set_column_widths(table, [1.5, 5.0, 1.5, 1.5, 3.0])


# === Hilfsfunktionen für Formatierung ===

def _add_section_heading(doc: Document, number: str, title: str):
    """Fügt eine Abschnittsüberschrift hinzu."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)

    # Nummer
    run = p.add_run(f"{number}  ")
    run.font.size = Pt(22)
    run.font.color.rgb = COLORS["secondary"]
    run.font.name = 'Calibri'

    # Titel
    run = p.add_run(title)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = COLORS["primary"]
    run.font.name = 'Calibri'


def _add_section_break(doc: Document):
    """Fügt einen Seitenumbruch hinzu."""
    doc.add_page_break()


def _add_thin_line(doc: Document):
    """Fügt eine dünne Trennlinie hinzu."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("─" * 80)
    run.font.size = Pt(6)
    run.font.color.rgb = COLORS["border"]
    run.font.name = 'Calibri'


def _style_header_cell(cell):
    """Formatiert eine Header-Zelle."""
    _set_cell_bg(cell, "1D1D1F")
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = COLORS["white"]
            run.font.name = 'Calibri'
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Padding reduzieren
    _set_cell_margins(cell, top=40, bottom=40, start=40, end=40)


def _style_body_cell(cell, bold=False, align="left", font_size=9, color=None):
    """Formatiert eine Body-Zelle."""
    for p in cell.paragraphs:
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in p.runs:
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run.font.name = 'Calibri'
            if color:
                run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    _set_cell_margins(cell, top=30, bottom=30, start=40, end=40)


def _set_cell_bg(cell, color_hex: str):
    """Setzt die Hintergrundfarbe einer Zelle."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _shade_row(row, color_hex: str):
    """Färbt eine komplette Zeile ein."""
    for cell in row.cells:
        _set_cell_bg(cell, color_hex)


def _set_cell_margins(cell, top=0, bottom=0, start=0, end=0):
    """Setzt die Ränder einer Zelle (in Twips)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}/>')
        tcPr.append(tcMar)
    for side, value in [('top', top), ('bottom', bottom), ('start', start), ('end', end)]:
        el = tcMar.find(qn(f'w:{side}'))
        if el is None:
            el = parse_xml(f'<w:{side} {nsdecls("w")} w:w="{value}" w:type="dxa"/>')
            tcMar.append(el)
        else:
            el.set(qn('w:w'), str(value))
            el.set(qn('w:type'), 'dxa')


def _set_column_widths(table, widths_cm: list):
    """Setzt die Spaltenbreiten einer Tabelle in cm."""
    for row in table.rows:
        for i, width in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(width)


def _get_free_label(dienst: Dienst) -> str:
    """Gibt das Label für einen freien Tag zurück."""
    free_event = next((e for e in dienst.events if e.is_free), None)
    if free_event:
        if free_event.dienst_type == DienstType.URLAUB:
            return "Urlaub"
        if free_event.dienst_type == DienstType.REISEZEITAUSGLEICH:
            return "RZA"
    return "Frei"
