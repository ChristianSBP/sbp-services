"""Tests für den Word-Writer (Docx-Ausgabe)."""

import pytest
from datetime import date, time
from pathlib import Path
import tempfile

from dienstplan.models.events import Event, Dienst, DienstType
from dienstplan.models.plan import Dienstplan
from dienstplan.output.word_writer import write_dienstplan_docx


def _make_test_plan():
    """Erstellt einen minimalen Testplan mit 2 Wochen."""
    start = date(2026, 3, 2)
    end = date(2026, 3, 15)
    dienste = []

    for i in range(14):
        d = date.fromordinal(start.toordinal() + i)
        if d.weekday() < 5:  # Mo-Fr
            dienste.append(Dienst(
                dienst_date=d,
                events=[Event(
                    event_date=d, start_time=time(10, 0), end_time=time(13, 0),
                    dienst_type=DienstType.PROBE, raw_text=f"Probe {d}",
                    programm="Beethoven: Sinfonie Nr. 5",
                    leitung="Kahle",
                    ort="GWH Bad Lausick",
                )],
                dienst_count=1.0,
            ))
        else:
            dienste.append(Dienst(dienst_date=d, dienst_count=0.0, is_free=True))

    return Dienstplan.from_events([], dienste, start, end)


class TestWordWriter:
    """Tests für die DOCX-Erzeugung."""

    def test_creates_docx_file(self, config):
        """DOCX-Datei wird erstellt."""
        plan = _make_test_plan()
        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / "test_dienstplan.docx"
            write_dienstplan_docx(plan, output, config)
            assert output.exists()
            assert output.stat().st_size > 1000  # Nicht leer

    def test_docx_has_content(self, config):
        """DOCX enthält erkennbare Tabellen."""
        from docx import Document

        plan = _make_test_plan()
        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / "test_content.docx"
            write_dienstplan_docx(plan, output, config)

            doc = Document(str(output))
            # Mindestens 2 Tabellen (Kalender + Tagesansicht)
            assert len(doc.tables) >= 2

    def test_docx_has_paragraphs(self, config):
        """DOCX enthält Überschriften."""
        from docx import Document

        plan = _make_test_plan()
        with tempfile.TemporaryDirectory() as tmpdir:
            output = Path(tmpdir) / "test_paragraphs.docx"
            write_dienstplan_docx(plan, output, config)

            doc = Document(str(output))
            texts = [p.text for p in doc.paragraphs]
            # Prüfe dass Schlüsseltexte vorhanden sind
            all_text = " ".join(texts).upper()
            assert "DIENSTPLAN" in all_text
            assert "BLÄSERPHILHARMONIE" in all_text

    def test_docx_config_formats(self, config, tvk_only_config):
        """DOCX funktioniert sowohl mit HTV als auch TVK Config."""
        plan = _make_test_plan()
        with tempfile.TemporaryDirectory() as tmpdir:
            # HTV
            htv_output = Path(tmpdir) / "test_htv.docx"
            write_dienstplan_docx(plan, htv_output, config)
            assert htv_output.exists()

            # TVK
            tvk_output = Path(tmpdir) / "test_tvk.docx"
            write_dienstplan_docx(plan, tvk_output, tvk_only_config)
            assert tvk_output.exists()
